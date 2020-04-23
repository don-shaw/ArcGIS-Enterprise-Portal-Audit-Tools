from arcgis.gis import GIS, RoleManager
from datetime import datetime
import pandas as pd
import csv
import docx
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION
import matplotlib.pyplot as plt
import seaborn as sns
from os import path
import os
import configparser
import logging
import keyring
import subprocess
import win32com.client
import warnings
warnings.filterwarnings('ignore')

plt.style.use('fivethirtyeight')


def create_directories(report_dir, today_dir):
    os.chdir(report_dir)
    if path.isdir(today_dir) is False:
        logging.info('Creating the Report Directory...')
        os.mkdir(path.join(report_dir, today_dir))
        os.chdir(today_dir)
        os.mkdir('csv_files')
        os.mkdir('figures')


def generate_sys_log_report(slp_directory, today_dir, server_log_dir):
    os.chdir(today_dir)
    if path.isdir(today_dir) is True:
        os.mkdir('sys_log_report')
        os.chdir(slp_directory)
        output_dir = path.join(today_dir, 'sys_log_report')
        cmd = (str(f"slp.exe -f AGSFS -i {server_log_dir} -d {output_dir} -eh now -sh 1080 -a complete -r spreadsheet -sbu true -o false, shell=True"))
        logging.info('Creating the System Log Report...')
        subprocess.call(cmd)


def get_portal_data(url, cred_name, portal_username, today_dir):

    try:
        # Get credentials
        portal_password = keyring.get_password(cred_name, portal_username)
        portal = GIS(url, portal_username, portal_password, verify_cert=False)
        logging.info('Querying the Enterprise Portal...')

        # Query users, groups, and items
        users = portal.users.search('!username:esri_*')
        groups = portal.groups.search('!owner:esri_*')
        all_items = portal.content.search(query='!owner:esri*', max_items=10000)

        # Create empty dictionaries, will be used to populate CSV files
        user_dict = {}
        group_dict = {}
        item_dict = {}

        # Get users and write them to CSV
        with open(path.join(today_dir, 'csv_files', 'user.csv'), 'w', newline='', encoding='utf-8') as user_csv:
            user_file = csv.DictWriter(user_csv,
                                       fieldnames=['USERNAME', 'EMAIL', 'ROLE', 'LAST LOGIN',
                                                   'CREATED', 'GROUPS', 'ITEMS'])
            user_file.writeheader()

            rm = RoleManager(portal)
            roles = rm.all()

            for user in users:
                # if user.idpUsername is not None:
                num_items = 0

                user_dict['USERNAME'] = user.username
                user_dict['EMAIL'] = user.email
                user_dict['ROLE'] = user.role
                for role in roles:
                    role_id = role.role_id
                    if role_id == user.roleId:
                        user_dict['ROLE'] = role.name

                if user.lastLogin != -1:
                    user_dict['LAST LOGIN'] = datetime.fromtimestamp(float(user.lastLogin / 1000)).strftime('%m/%d/%Y')
                else:
                    user_dict['LAST LOGIN'] = -1
                user_dict['CREATED'] = datetime.fromtimestamp(float(user.created / 1000)).strftime('%m/%d/%Y')

                user_groups = user.groups
                g_list = []

                for g in user_groups:
                    g_list.append(g.title)
                user_dict['GROUPS'] = str(g_list)[1:-1]

                user_content = user.items()
                folders = user.folders

                for item in user_content:
                    num_items += 1

                for folder in folders:
                    folder_items = user.items(folder=folder['title'])
                    for item in folder_items:
                        num_items += 1
                user_dict['ITEMS'] = num_items
                user_file.writerow(user_dict)
        logging.info('User File:    {0}'.format(path.join(today_dir, 'csv_files', 'user.csv')))

        # Get Groups
        with open(path.join(today_dir, 'csv_files', 'groups.csv'), 'w', newline='', encoding='utf-8') as group_csv:
            groups_file = csv.DictWriter(group_csv, fieldnames=['TITLE', 'OWNER', 'MANAGERS', 'USERS', 'ITEMS'])
            groups_file.writeheader()

            for g in groups:

                members = g.get_members()

                group_dict['TITLE'] = g.title
                group_dict['OWNER'] = members['owner']
                group_dict['MANAGERS'] = str(str(members['admins']).replace("'", ''))[1:-1]
                group_dict['USERS'] = str(str(members['users']).replace("'", ''))[1:-1]
                group_dict['ITEMS'] = len(g.content())
                groups_file.writerow(group_dict)
        logging.info('Group File:    {0}'.format(path.join(today_dir, 'csv_files', 'groups.csv')))

        # Get Items
        with open(path.join(today_dir, 'csv_files', 'items.csv'), 'w', newline='', encoding='utf-8') as items_csv:
            items_file = csv.DictWriter(items_csv, fieldnames=['TITLE', 'OWNER', 'ID', 'TYPE', 'AUTHORITATIVE',
                                                               'TAGS', 'ACCESS', 'SHARED WITH ORG',
                                                               'SHARED WITH EVERYONE', 'SHARED WITH GROUPS', 'VIEWS',
                                                               'CREATED', 'HOMEPAGE', 'THUMBNAIL', 'DESCRIPTION', 'SIZE'])

            items_file.writeheader()
            for item in all_items:
                if item.type in ['Geoprocessing Service', 'Service Definition', 'Code Attachment', 'Geometry Service',
                                 'Vector Tile Service', 'Vector Tile Package']:
                    pass
                else:
                    print(item)
                    item_groups = []
                    item_dict['TITLE'] = item.title
                    item_dict['OWNER'] = item.owner
                    item_dict['ID'] = item.id
                    item_dict['TYPE'] = item.type
                    item_dict['AUTHORITATIVE'] = item.content_status
                    item_dict['TAGS'] = str(item.tags)[1:-1]
                    item_dict['DESCRIPTION'] = item.description
                    item_dict['VIEWS'] = item.numViews
                    item_dict['CREATED'] = datetime.fromtimestamp(float(item.created / 1000)).strftime('%m/%d/%Y')
                    # item_dict['dependent_on'] = item.dependent_upon()
                    # item_dict['dependent_to'] = item.dependent_to()
                    item_dict['HOMEPAGE'] = item.homepage
                    item_dict['SHARED WITH EVERYONE'] = item.shared_with['everyone']
                    item_dict['SHARED WITH ORG'] = item.shared_with['org']
                    for g in item.shared_with['groups']:
                        item_groups.append(g.title)
                    item_dict['SHARED WITH GROUPS'] = str(item_groups)[1:-1]
                    # print(item_groups)
                    item_dict['ACCESS'] = item.access
                    item_dict['SIZE'] = item.size / 1000 / 1000
                    item_dict['THUMBNAIL'] = item.thumbnail
                    items_file.writerow(item_dict)

        logging.info('Item File:    {0}'.format(path.join(today_dir, 'csv_files', 'items.csv')))
        process_plots(today_dir)
        return groups
    except Exception as e:
        logging.error(e)


def process_plots(today_dir):
    # Create the dataframes used for the plots
    items_df = pd.read_csv(path.join(today_dir, 'csv_files', 'items.csv'))
    items_df = items_df[['TITLE', 'OWNER', 'TYPE', 'AUTHORITATIVE', 'TAGS', 'ACCESS',
                        'SHARED WITH GROUPS', 'VIEWS', 'CREATED']]
    group_df = pd.read_csv(path.join(today_dir, 'csv_files', 'groups.csv'))

    logging.info('Processing Plots...')

    # Make top 5 web maps
    top_5_web_maps = items_df.loc[items_df['TYPE'] == 'Web Map'].sort_values('VIEWS', ascending=False).head(5)

    plt.figure(figsize=(10, 10))
    sns.barplot(x='TITLE', y='VIEWS', palette='Set1', data=top_5_web_maps)
    plt.title('Five Most Viewed Web Maps: {0}'.format(datetime.now().strftime("%m/%d/%Y")))
    plt.xticks(rotation=90)
    plt.tight_layout()
    plt.savefig(path.join(today_dir, 'figures', '{0}_top_5_web_maps'.format(datetime.now().strftime("%m%d%Y"))))

    # Make top 5 web apps

    top_5_web_apps = items_df.loc[items_df['TYPE'] == 'Web Mapping Application'].sort_values('VIEWS', ascending=False).head(5)
    plt.figure(figsize=(10, 10))

    sns.barplot(x='TITLE', y='VIEWS', palette='Set1', data=top_5_web_apps)
    plt.title('Five Most Viewed Web Apps: {0}'.format(datetime.now().strftime("%m/%d/%Y")))
    plt.xticks(rotation=90)
    plt.tight_layout()
    plt.savefig(path.join(today_dir, 'figures', '{0}_top_5_web_apps'.format(datetime.now().strftime("%m%d%Y"))))


    # Make top 5 groups

    plt.figure(figsize=(10, 10))

    sns.barplot(x='TITLE', y='ITEMS', palette='Set1', data=group_df.sort_values('ITEMS', ascending=False))
    plt.title('Groups With the Most Content: {0}'.format(datetime.now().strftime("%m/%d/%Y")))
    plt.xticks(rotation=90)
    plt.tight_layout()
    plt.savefig(path.join(today_dir, 'figures', '{0}_top_groups'.format(datetime.now().strftime("%m%d%Y"))))

def requests_last_two_weeks(throughput, today_dir):
    plt.figure(figsize=(20,10))
    plt.style.use('fivethirtyeight')

    by_day = throughput.groupby('date')['HTTP 200'].sum().reset_index()
    by_day = throughput[throughput['date'] > pd.Timedelta(-14, unit='d') + pd.datetime.today().date()]


    sns.barplot(x='date', y='HTTP 200', color='blue', ci=False, alpha=.8, data=by_day)
    plt.xticks(rotation=90)
    plt.xlabel('Date', fontsize=20, fontweight='bold')
    plt.ylabel('Requests', fontsize=20, fontweight='bold')
    plt.xticks(fontsize=13)
    plt.title('Total Number of Requests in the last 14 days', fontweight='bold')
    plt.tight_layout()
    plt.savefig(path.join(today_dir, 'figures', 'Requests in the last 14 days.png'))

def most_popular_items_plot(most_popular_resources, today_dir):
    plt.figure(figsize=(8,8))
    plt.style.use('fivethirtyeight')

    sns.barplot(x='Resource', y='Requests', ci=None, data = most_popular_resources.sort_values(by='Requests', ascending=False))
    plt.xticks(rotation=90)
    plt.xlabel('Resource', fontsize=20, fontweight='bold')
    plt.ylabel('Requests', fontsize=20, fontweight='bold')
    plt.xticks(fontsize=13)
    plt.title('5 Most Popular Services in the Last 14 Days', fontweight='bold')
    plt.tight_layout()
    plt.legend('')
    plt.savefig(path.join(today_dir, 'figures', 'Most Popular Resources.png'))


def process_sys_log_report(today_dir):
    report_dir = path.join(today_dir, 'sys_log_report')
    os.chdir(report_dir)
    for file in os.listdir(report_dir):
        if file.endswith('xlsx'):
            report = file
    # System Log Parser dfs
    stats_by_user = pd.read_excel(report, sheet_name='Statistics By User', header=4)
    stats_by_resource = pd.read_excel(report, sheet_name='Statistics By Resource', header=4)
    resources_by_time = pd.read_excel(report, sheet_name='Elapsed Time - All Resources', header=3)
    throughput = pd.read_excel(report, sheet_name='Throughput per Minute', header=3)
    throughput['date'] = pd.to_datetime(throughput['Date Time (Local Time)']).dt.to_period('D')
    stats_by_user = stats_by_user[stats_by_user['User'].str.len() < 15]
    resources_by_time = resources_by_time[resources_by_time['User'].str.len() < 15]
    by_user = stats_by_user.groupby('User')['Count'].sum().reset_index()
    by_user = by_user[by_user['User'] != '-']
    most_active_users = by_user.head(10)
    most_active_users['Requests'] = most_active_users['Count']
    most_active_users = most_active_users.drop(columns='Count').sort_values(by='Requests', ascending=False)
    items_df = pd.read_csv(path.join(today_dir, 'csv_files', 'items.csv'))
    items_df = items_df[items_df['TYPE'].isin(['Map Service', 'Feature Service'])]
    items_df.loc[items_df['TYPE'] == 'Map Service', 'Resource'] = items_df['TITLE'] + '.MapServer'
    items_df.loc[items_df['TYPE'] == 'Feature Service', 'Resource'] = items_df['TITLE'] + '.FeatureServer'
    last_accessed = resources_by_time.groupby('Resource')['Date Time (Local Time)'].max().reset_index()
    last_accessed['LAST ACCESSED'] = pd.to_datetime(last_accessed['Date Time (Local Time)'])
    last_accessed = last_accessed[last_accessed['Resource'].str.contains('GPServer') == False]
    item_metrics = items_df.merge(right=last_accessed, how='outer', left_on='Resource', right_on='Resource')
    has_not_been_accessed = item_metrics[
        item_metrics['LAST ACCESSED'] < pd.datetime.today() - pd.Timedelta(14, unit='d')]
    has_not_been_accessed = has_not_been_accessed[
        ['TITLE', 'OWNER', 'TYPE', 'ACCESS', 'VIEWS', 'CREATED', 'LAST ACCESSED']]
    has_not_been_accessed = has_not_been_accessed.dropna()
    resources_by_time['date'] = pd.to_datetime(resources_by_time['Date Time (Local Time)'])
    resource_by_day_total_requests = resources_by_time.groupby(['Resource', 'date'])['HTTP Code'].agg(
        'count').reset_index()
    resource_by_day_total_requests = resource_by_day_total_requests[resource_by_day_total_requests['Resource'].isin(
        ['System/PublishingTools.GPServer', 'System/PublishingToolsEx.GPServer',
         'Utilities/PrintingTools.GPServer']) == False]
    resource_by_day_total_requests = resource_by_day_total_requests[
        resource_by_day_total_requests['date'] > pd.Timedelta(-14, unit='d') + datetime.today()]
    most_popular_resources = resource_by_day_total_requests.groupby('Resource')['HTTP Code'].sum().reset_index().head(5)
    most_popular_resources['Requests'] = most_popular_resources['HTTP Code']
    most_popular_resources = most_popular_resources.drop(columns='HTTP Code').sort_values(by='Requests',
                                                                                          ascending=False).reset_index()
    requests_last_two_weeks(throughput, today_dir)
    most_popular_items_plot(most_popular_resources, today_dir)

    return has_not_been_accessed, most_popular_resources, most_active_users


def process_report(url, cred_name, portal_username, template, today_dir):
    try:

        results = get_portal_data(url, cred_name, portal_username, today_dir)
        has_not_been_accessed, most_popular_resources, most_active_users = process_sys_log_report(today_dir)

        logging.info('Processing Report...')
        # Get dataframes
        user_df = pd.read_csv(path.join(today_dir, 'csv_files', 'user.csv'))
        group_df = pd.read_csv(path.join(today_dir, 'csv_files', 'groups.csv'))
        items_df = pd.read_csv(path.join(today_dir, 'csv_files', 'items.csv'))
        items_df = items_df[['TITLE', 'OWNER', 'TYPE', 'AUTHORITATIVE', 'TAGS', 'ACCESS',
                             'SHARED WITH GROUPS', 'VIEWS', 'CREATED']]
        items_no_tags_df = items_df[items_df['TAGS'].isnull()]

        items_df['created_date'] = pd.to_datetime(items_df['CREATED'])
        items_created_last_two_weeks = items_df[items_df['created_date'] > pd.Timedelta(-14, unit='d') + datetime.today()]
        items_created_last_two_weeks = items_created_last_two_weeks.drop(columns='created_date')

        # Make document
        doc = docx.Document(template)
        # title = doc.add_heading(f'Enterprise Community Portal Audit', 0)
        # title.alignment = 1
        # date = doc.add_paragraph(f'{datetime.now().strftime("%m/%d/%Y")}')
        # date.alignment = 1
        #
        # doc.add_paragraph('')
        # doc.add_paragraph('')

        # Make User Table
        doc.add_heading('Enterprise Portal Users', level=1)
        user_table = doc.add_table(user_df.shape[0] + 1, user_df.shape[1], style='Light Grid Accent 5')
        user_table.autofit = False

        for j in range(user_df.shape[-1]):
            user_table.cell(0, j).text = user_df.columns[j]

        for i in range(user_df.shape[0]):
            for j in range(user_df.shape[-1]):
                user_table.cell(i + 1, j).text = str(user_df.values[i, j])

        doc.add_page_break()

        # Make Group Table
        doc.add_heading('Enterprise Groups', level=1)
        group_table = doc.add_table(group_df.shape[0] + 1, group_df.shape[1], style='Light Grid Accent 5')
        group_table.autofit = False
        for j in range(group_df.shape[-1]):
            group_table.cell(0, j).text = group_df.columns[j]

        for i in range(group_df.shape[0]):
            for j in range(group_df.shape[-1]):
                group_table.cell(i + 1, j).text = str(group_df.values[i, j])
        doc.add_paragraph('')

        # List Each Group
        # Make each group
        doc.add_heading('Items by Group', level=1)
        for g in results:
            if len(g.content()) > 0:
                doc.add_heading(f'{g.title}', level=2)
                content = g.content()
                gdf = pd.DataFrame(content)

                gdf['authoritative'] = gdf['contentStatus']
                gdf['views'] = gdf['numViews']
                gdf = gdf[['title', 'owner', 'type', 'access', 'authoritative', 'views', 'tags', 'thumbnail']]
                g_table = doc.add_table(gdf.shape[0] + 1, gdf.shape[1], style='Light Grid Accent 5')
                g_table.autofit = False
                for j in range(gdf.shape[-1]):
                    g_table.cell(0, j).text = gdf.columns[j]

                for i in range(gdf.shape[0]):
                    for j in range(gdf.shape[-1]):
                        g_table.cell(i + 1, j).text = str(gdf.values[i, j])
            else:
                pass
            doc.add_paragraph('')
        # Make Item Table
        # current_section = doc.sections[-1]
        # new_width, new_height = current_section.page_height, current_section.page_width
        # new_section = doc.add_section(WD_SECTION.NEW_PAGE)
        # new_section.orientation = WD_ORIENT.LANDSCAPE
        # new_section.page_width = new_width
        # new_section.page_height = new_height
        # new_section.left_margin = Inches(0.25)
        # new_section.right_margin = Inches(0.25)

        doc.add_heading('Portal Items Created in the Last 14 Days', level=1)
        item_table = doc.add_table(items_created_last_two_weeks.shape[0] + 1, items_created_last_two_weeks.shape[1], style='Light Grid Accent 5')
        item_table.autofit = False
        for j in range(items_created_last_two_weeks.shape[-1]):
            item_table.cell(0, j).text = items_created_last_two_weeks.columns[j]

        for i in range(items_created_last_two_weeks.shape[0]):
            for j in range(items_created_last_two_weeks.shape[-1]):
                item_table.cell(i + 1, j).text = str(items_created_last_two_weeks.values[i, j])

        doc.add_page_break()

        # Items with no tags
        doc.add_heading('Items With No Tags', level=1)
        no_tag_table = doc.add_table(items_no_tags_df.shape[0] + 1,
                                     items_no_tags_df.shape[1], style='Light Grid Accent 5')
        no_tag_table.autofit = False
        for j in range(items_no_tags_df.shape[-1]):
            no_tag_table.cell(0, j).text = items_no_tags_df.columns[j]

        for i in range(items_no_tags_df.shape[0]):
            for j in range(items_no_tags_df.shape[-1]):
                no_tag_table.cell(i + 1, j).text = str(items_no_tags_df.values[i, j])
        doc.add_paragraph('')

        # Inactive Items
        doc.add_heading('Items Not Accessed in Last 14 Days', level=1)
        not_accessed_table = doc.add_table(has_not_been_accessed.shape[0] + 1,
                                     has_not_been_accessed.shape[1], style='Light Grid Accent 5')
        not_accessed_table.autofit = False
        for j in range(has_not_been_accessed.shape[-1]):
            not_accessed_table.cell(0, j).text = has_not_been_accessed.columns[j]

        for i in range(has_not_been_accessed.shape[0]):
            for j in range(has_not_been_accessed.shape[-1]):
                not_accessed_table.cell(i + 1, j).text = str(has_not_been_accessed.values[i, j])
        doc.add_paragraph('')

        # Most Popular Resources
        doc.add_heading('Most Popular Resources in Last 14 Days', level=1)
        most_popular_table = doc.add_table(most_popular_resources.shape[0] + 1,
                                           most_popular_resources.shape[1], style='Light Grid Accent 5')
        most_popular_table.autofit = False
        for j in range(most_popular_resources.shape[-1]):
            most_popular_table.cell(0, j).text = most_popular_resources.columns[j]

        for i in range(most_popular_resources.shape[0]):
            for j in range(most_popular_resources.shape[-1]):
                most_popular_table.cell(i + 1, j).text = str(most_popular_resources.values[i, j])
        doc.add_paragraph('')

        # Most Popular Resources
        doc.add_heading('Most Active Users in Last 14 Days', level=1)
        active_user_table = doc.add_table(most_active_users.shape[0] + 1,
                                           most_active_users.shape[1], style='Light Grid Accent 5')
        active_user_table.autofit = False
        for j in range(most_active_users.shape[-1]):
            active_user_table.cell(0, j).text = most_active_users.columns[j]

        for i in range(most_active_users.shape[0]):
            for j in range(most_active_users.shape[-1]):
                active_user_table.cell(i + 1, j).text = str(most_active_users.values[i, j])
        doc.add_paragraph('')

        # Start adding pictures

        # current_section = doc.sections[-1]
        # new_width, new_height = current_section.page_height, current_section.page_width
        # new_section = doc.add_section(WD_SECTION.NEW_PAGE)
        # new_section.orientation = WD_ORIENT.PORTRAIT
        # new_section.page_width = new_width
        # new_section.page_height = new_height

        doc.add_heading('Top 5 Web Maps', level=1)
        doc.add_picture(path.join(today_dir, 'Figures', '{0}_top_5_web_maps.png'.format(datetime.now().strftime("%m%d%Y"))),
                        width=Inches(7))

        doc.add_heading('Top 5 Web Apps', level=1)
        doc.add_picture(path.join(today_dir, 'Figures', '{0}_top_5_web_apps.png'.format(datetime.now().strftime("%m%d%Y"))),
                        width=Inches(7))

        doc.add_heading('Top Groups by Content', level=1)
        doc.add_picture(path.join(today_dir, 'Figures', '{0}_top_groups.png'.format(datetime.now().strftime("%m%d%Y"))),
                        width=Inches(7))

        doc.add_heading('Most Popular Resources', level=1)
        doc.add_picture(path.join(today_dir, 'Figures', 'Most Popular Resources.png'),
                        width=Inches(7))

        current_section = doc.sections[-1]
        new_width, new_height = current_section.page_height, current_section.page_width
        new_section = doc.add_section(WD_SECTION.NEW_PAGE)
        new_section.orientation = WD_ORIENT.LANDSCAPE
        new_section.page_width = new_width
        new_section.page_height = new_height
        new_section.left_margin = Inches(0.25)
        new_section.right_margin = Inches(0.25)

        doc.add_heading('Portal Traffic in the Last 14 Days', level=1)
        doc.add_picture(path.join(today_dir, 'Figures', 'Requests in the last 14 days.png'),
                        width=Inches(10), height=Inches(5.5))

        # save the doc
        output_file = path.join(today_dir, '{0}_Audit.docx'.format(datetime.now().strftime("%m%d%Y")))
        doc.save(output_file)

        word = win32com.client.DispatchEx("Word.Application")
        doc2 = word.Documents.Open(output_file)
        doc2.TablesOfContents(1).Update()
        doc2.Close(SaveChanges=True)
        word.Quit()
        logging.info('Report File:    {0}\n'.format(output_file))

    except Exception as e:
        logging.error(e)
        print(e)


if __name__ == '__main__':

    log_dir = path.dirname(path.realpath(__file__))
    log_name = path.join(log_dir, 'portal_audit_tools.log')
    log_format = "%(asctime)s - %(levelname)s - %(message)s"
    logging.basicConfig(filename=log_name,
                        level=logging.INFO,
                        format=log_format,
                        filemode="w")
    logging.getLogger('arcgis').setLevel(logging.WARNING)

    config = configparser.ConfigParser()
    config.read(path.join(log_dir, 'config.ini'))

    portal_url = config.get('ALL', 'portal_url')
    portal_cred_name = config.get('ALL', 'portal_cred_name')
    portal_cred_user = config.get('ALL', 'portal_cred_user')
    reports_directory = config.get('ALL', 'reports_directory')
    today_directory = path.join(reports_directory, f'{datetime.now().strftime("%m-%d-%Y")}')
    report_template = config.get('ALL', 'report_template')
    system_log_parser = config.get('ALL', 'sys_log_directory')
    server_log_directory = config.get('ALL', 'server_log_directory')

    logging.info("***** Start time:  {0}\n".format(datetime.now().strftime("%A %B %d %I:%M:%S %p %Y")))
    logging.info('Portal URL:   {0}'.format(portal_url))
    logging.info('Windows Credential:   {0}'.format(portal_cred_name))
    logging.info('Portal Username:  {0}'.format(portal_cred_user))
    logging.info('Audit Report Directory:   {0}\n'.format(today_directory))

    # Go!
    try:
        create_directories(reports_directory, today_directory)
        generate_sys_log_report(system_log_parser, today_directory, server_log_directory)
        process_report(portal_url, portal_cred_name, portal_cred_user, report_template, today_directory)
    except Exception as e:
        print(e)
    logging.info("***** Completed time:  {0}\n".format(datetime.now().strftime("%A %B %d %I:%M:%S %p %Y")))
    logging.shutdown()
